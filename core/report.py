"""
报告生成模块
支持 HTML（内置）/ PDF（reportlab）/ DOCX（python-docx）三种格式。

使用方式：
    from core.report import ReportConfig, ReportAssets, collect_map_screenshot,
                             collect_viz_figures, export_html, export_pdf, export_docx
"""

import base64
import io
import os
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional


# ─────────────────────────────────────────────────────────────────────────────
#  数据结构
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class ReportConfig:
    title:           str  = "海冰多参数综合反演分析报告"
    author:          str  = ""
    note:            str  = ""
    include_map:     bool = True
    include_figures: bool = True
    include_meta:    bool = True
    fmt_html:        bool = True
    fmt_pdf:         bool = False
    fmt_docx:        bool = False
    out_dir:         str  = ""


@dataclass
class ReportAssets:
    timestamp: str        = ""
    map_png:   Optional[bytes] = None
    figures:   list       = field(default_factory=list)  # [{"title": str, "png": bytes}]
    metadata:  dict       = field(default_factory=dict)


# ─────────────────────────────────────────────────────────────────────────────
#  数据收集
# ─────────────────────────────────────────────────────────────────────────────

def collect_map_screenshot(map_widget) -> Optional[bytes]:
    """
    截取地图 Widget 为 PNG bytes。

    QWebEngineView 由独立 GPU 进程渲染，widget.grab() 只能拿到灰色占位层。
    改用 QScreen.grabWindow(native_winId) 在操作系统合成层截图。
    调用前应确保地图 Tab 已激活并处理完渲染事件。
    """
    try:
        from PyQt6.QtCore import QBuffer, QByteArray, QEventLoop, QIODevice, QTimer
        from PyQt6.QtWidgets import QApplication
        from PyQt6.QtWebEngineWidgets import QWebEngineView

        # 找到内部 QWebEngineView（若 map_widget 本身就是则直接用）
        if isinstance(map_widget, QWebEngineView):
            view = map_widget
        else:
            view = map_widget.findChild(QWebEngineView)

        # 等待 OS 合成完成（对话框 hide 后需要一帧让系统重绘）
        loop = QEventLoop()
        QTimer.singleShot(500, loop.quit)
        loop.exec()
        QApplication.processEvents()

        # 用 OS 合成层截图：grabWindow(winId) 能抓到 GPU 渲染内容
        screen = QApplication.primaryScreen()
        if view and view.winId():
            px = screen.grabWindow(int(view.winId()))
        else:
            # 回退：用 mapToGlobal 定位 widget 在屏幕上的坐标后区域截图
            from PyQt6.QtCore import QPoint
            target = view if view else map_widget
            gpos = target.mapToGlobal(QPoint(0, 0))
            px = screen.grabWindow(
                0, gpos.x(), gpos.y(), target.width(), target.height()
            )

        if px.isNull():
            raise RuntimeError("grabWindow 返回空 Pixmap")

        buf = QByteArray()
        qbuf = QBuffer(buf)
        qbuf.open(QIODevice.OpenModeFlag.WriteOnly)
        px.save(qbuf, "PNG")
        qbuf.close()
        return bytes(buf)
    except Exception as e:
        print(f"[报告] 地图截图失败: {e}")
        return None


def collect_viz_figures(plugin_manager) -> list:
    """从所有已加载插件的可视化 Tab 中抓取 matplotlib 图表（PNG bytes）。"""
    figures = []
    try:
        from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg
    except ImportError:
        return figures

    for pid in plugin_manager.get_loaded_ids():
        plugin = plugin_manager.get_plugin(pid)
        if not plugin:
            continue
        try:
            for tab_name, widget in plugin.get_viz_tabs():
                canvas = widget.findChild(FigureCanvasQTAgg)
                if canvas and canvas.figure:
                    buf = io.BytesIO()
                    canvas.figure.savefig(
                        buf, format="png", dpi=130,
                        facecolor=canvas.figure.get_facecolor(),
                        bbox_inches="tight",
                    )
                    figures.append({
                        "title": f"{plugin.name} · {tab_name}",
                        "png":   buf.getvalue(),
                    })
        except Exception as e:
            print(f"[报告] 图表抓取失败 {pid}: {e}")
    return figures


# ─────────────────────────────────────────────────────────────────────────────
#  内部工具：CJK 字体注册
# ─────────────────────────────────────────────────────────────────────────────

def _find_cjk_font_path() -> tuple:
    """
    返回 (font_file_path, is_ttc)。
    优先用 matplotlib 的字体管理器定位系统已有的 CJK 字体；
    若 matplotlib 不可用则枚举常见路径。
    """
    # ── 方案 A：借助 matplotlib font_manager（已在 seaice 环境中安装）
    try:
        import matplotlib.font_manager as fm
        _cjk_names = [
            "PingFang SC", "PingFang TC", "Heiti SC", "STHeiti",
            "Microsoft YaHei", "SimHei", "Arial Unicode MS",
            "Noto Sans CJK SC", "WenQuanYi Micro Hei",
        ]
        for name in _cjk_names:
            try:
                path = fm.findfont(
                    fm.FontProperties(family=name),
                    fallback_to_default=False,
                )
                # findfont 找不到时会返回 DejaVuSans 路径，排除
                if path and "DejaVu" not in path and os.path.exists(path):
                    is_ttc = path.lower().endswith(".ttc")
                    return path, is_ttc
            except Exception:
                continue
    except ImportError:
        pass

    # ── 方案 B：固定路径枚举
    _candidates = [
        ("/Library/Fonts/Arial Unicode MS.ttf",                       False),
        ("/System/Library/Fonts/Supplemental/Arial Unicode MS.ttf",   False),
        ("/System/Library/Fonts/PingFang.ttc",                        True),
        ("/System/Library/Fonts/STHeiti Light.ttc",                   True),
        ("C:/Windows/Fonts/msyh.ttf",                                 False),
        ("C:/Windows/Fonts/simsun.ttc",                               True),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",    True),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",    True),
    ]
    for path, is_ttc in _candidates:
        if os.path.exists(path):
            return path, is_ttc

    return "", False


def _register_cjk_font(pdfmetrics, TTFont) -> str:
    """
    向 reportlab 注册 CJK 字体，返回注册后的字体名称。
    若注册失败，回退到 Helvetica（文字会显示黑框，但至少不崩溃）。
    """
    # 已注册则直接复用
    try:
        pdfmetrics.getFont("CNFont")
        return "CNFont"
    except Exception:
        pass

    path, is_ttc = _find_cjk_font_path()
    if not path:
        print("[报告] 未找到 CJK 字体，PDF 中文可能显示异常")
        return "Helvetica"

    # TTC 文件需要指定子字体索引；尝试 index 0、1
    try:
        if is_ttc:
            pdfmetrics.registerFont(TTFont("CNFont", path, subfontIndex=0))
        else:
            pdfmetrics.registerFont(TTFont("CNFont", path))
        print(f"[报告] 已注册 CJK 字体: {os.path.basename(path)}")
        return "CNFont"
    except Exception as e:
        # TTC index 0 失败时尝试 index 1
        if is_ttc:
            try:
                pdfmetrics.registerFont(TTFont("CNFont", path, subfontIndex=1))
                print(f"[报告] 已注册 CJK 字体 (subfont=1): {os.path.basename(path)}")
                return "CNFont"
            except Exception:
                pass
        print(f"[报告] CJK 字体注册失败 ({path}): {e}")
        return "Helvetica"


# ─────────────────────────────────────────────────────────────────────────────
#  HTML 生成
# ─────────────────────────────────────────────────────────────────────────────

def _b64_png(data: bytes) -> str:
    return base64.b64encode(data).decode()


def build_html(config: ReportConfig, assets: ReportAssets) -> str:
    sections = []

    if config.include_meta:
        meta_rows = "".join(
            f"<tr><td>{k}</td><td>{v}</td></tr>"
            for k, v in assets.metadata.items()
        )
        sections.append(f"""
<h2>元数据信息</h2>
<table class="meta">
  <tr><td><b>报告时间</b></td><td>{assets.timestamp}</td></tr>
  <tr><td><b>作者</b></td><td>{config.author or "—"}</td></tr>
  {meta_rows}
</table>""")

    if config.include_map and assets.map_png:
        b64 = _b64_png(assets.map_png)
        sections.append(f"""
<h2>地图视图</h2>
<div class="figure">
  <img src="data:image/png;base64,{b64}" alt="地图截图" />
  <p class="caption">图：当前地图视图截图</p>
</div>""")

    if config.include_figures:
        for i, fig in enumerate(assets.figures, 1):
            b64 = _b64_png(fig["png"])
            title = fig["title"]
            sections.append(f"""
<h2>{title}</h2>
<div class="figure">
  <img src="data:image/png;base64,{b64}" alt="{title}" />
  <p class="caption">图 {i}：{title}</p>
</div>""")

    note_html = (
        f'<div class="note"><b>备注：</b>{config.note}</div>'
        if config.note else ""
    )

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{config.title}</title>
<style>
  body {{
    font-family: "PingFang SC", "Microsoft YaHei", "Noto Sans CJK SC", sans-serif;
    margin: 48px auto; max-width: 940px; color: #1a2a3a; line-height: 1.7;
  }}
  h1 {{
    text-align: center; color: #062040;
    border-bottom: 3px solid #00C4E8; padding-bottom: 14px;
    margin-bottom: 6px; font-size: 24px;
  }}
  .subtitle {{ text-align: center; color: #5580a0; font-size: 13px; margin-bottom: 32px; }}
  h2 {{
    color: #0a4070; border-left: 4px solid #00C4E8;
    padding-left: 12px; margin-top: 40px; font-size: 17px;
  }}
  .figure {{ text-align: center; margin: 20px 0; }}
  .figure img {{
    max-width: 100%; border: 1px solid #d0dde8;
    border-radius: 6px; box-shadow: 0 2px 10px rgba(0,0,0,.10);
  }}
  .caption {{ color: #6680a0; font-size: 12.5px; margin-top: 7px; }}
  table.meta {{
    border-collapse: collapse; width: 100%; margin: 14px 0;
    font-size: 13px;
  }}
  table.meta td {{ border: 1px solid #d8e4ee; padding: 8px 14px; }}
  table.meta tr:nth-child(odd) td {{ background: #f4f9ff; }}
  .note {{
    background: #fffbe6; border-left: 4px solid #f0c040;
    padding: 10px 18px; border-radius: 4px; margin: 24px 0; font-size: 13px;
  }}
  footer {{
    text-align: center; margin-top: 64px; color: #aabbcc; font-size: 12px;
    border-top: 1px solid #e8eef4; padding-top: 18px;
  }}
</style>
</head>
<body>
<h1>{config.title}</h1>
<p class="subtitle">生成时间：{assets.timestamp}{"　|　作者：" + config.author if config.author else ""}</p>
{note_html}
{"".join(sections)}
<footer>由极地海冰多参数综合反演平台自动生成 &nbsp;·&nbsp; {assets.timestamp}</footer>
</body>
</html>"""


def export_html(config: ReportConfig, assets: ReportAssets, path: str):
    html = build_html(config, assets)
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)


# ─────────────────────────────────────────────────────────────────────────────
#  PDF 生成（依赖 reportlab）
# ─────────────────────────────────────────────────────────────────────────────

def export_pdf(config: ReportConfig, assets: ReportAssets, path: str):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            HRFlowable, Image, Paragraph, SimpleDocTemplate, Spacer, Table,
            TableStyle,
        )
    except ImportError:
        raise RuntimeError("PDF 导出需要安装 reportlab：pip install reportlab")

    # 注册中文字体
    cn_font = _register_cjk_font(pdfmetrics, TTFont)

    W, _ = A4
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.2*cm, bottomMargin=2.2*cm,
    )

    def S(name, **kw):
        return ParagraphStyle(name, fontName=cn_font, **kw)

    s_title   = S("T", fontSize=20, alignment=1, spaceAfter=4,
                  textColor=colors.HexColor("#062040"), leading=26)
    s_sub     = S("Sub", fontSize=10, alignment=1, spaceAfter=20,
                  textColor=colors.HexColor("#5580a0"))
    s_h2      = S("H2", fontSize=13, spaceBefore=20, spaceAfter=8,
                  textColor=colors.HexColor("#0a4070"), leading=18)
    s_body    = S("B", fontSize=10, spaceAfter=6,
                  textColor=colors.HexColor("#1a2a3a"), leading=15)
    s_caption = S("Cap", fontSize=9, alignment=1, spaceAfter=14,
                  textColor=colors.HexColor("#6680a0"))
    s_note    = S("Note", fontSize=10, spaceAfter=10,
                  textColor=colors.HexColor("#6a5000"),
                  backColor=colors.HexColor("#fffbe6"),
                  borderPadding=(6, 10, 6, 10))

    story = []
    story.append(Paragraph(config.title, s_title))
    sub_txt = f"生成时间：{assets.timestamp}"
    if config.author:
        sub_txt += f"　|　作者：{config.author}"
    story.append(Paragraph(sub_txt, s_sub))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#00C4E8")))
    story.append(Spacer(1, 0.4*cm))

    if config.note:
        story.append(Paragraph(f"备注：{config.note}", s_note))

    # 元数据表
    if config.include_meta:
        story.append(Paragraph("元数据信息", s_h2))
        rows = [["报告时间", assets.timestamp], ["作者", config.author or "—"]]
        rows += [[str(k), str(v)] for k, v in assets.metadata.items()]
        tbl = Table(rows, colWidths=[4.5*cm, 11*cm])
        tbl.setStyle(TableStyle([
            ("FONTNAME",    (0, 0), (-1, -1), cn_font),
            ("FONTSIZE",    (0, 0), (-1, -1), 9),
            ("BACKGROUND",  (0, 0), (0, -1),  colors.HexColor("#eef5ff")),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1),
             [colors.white, colors.HexColor("#f4f9ff")]),
            ("GRID",        (0, 0), (-1, -1), 0.5, colors.HexColor("#c0cce0")),
            ("PADDING",     (0, 0), (-1, -1), 6),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.4*cm))

    max_w = W - 4.4*cm

    # 地图截图
    if config.include_map and assets.map_png:
        story.append(Paragraph("地图视图", s_h2))
        story.append(Image(io.BytesIO(assets.map_png), width=max_w, height=max_w * 0.58))
        story.append(Paragraph("图：当前地图视图截图", s_caption))

    # 插件图表
    if config.include_figures:
        for i, fig in enumerate(assets.figures, 1):
            story.append(Paragraph(fig["title"], s_h2))
            story.append(Image(io.BytesIO(fig["png"]), width=max_w, height=max_w * 0.65))
            story.append(Paragraph(f"图 {i}：{fig['title']}", s_caption))

    story.append(Spacer(1, 0.8*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d8e4ee")))
    story.append(Paragraph(
        f"由极地海冰多参数综合反演平台自动生成 · {assets.timestamp}", s_caption
    ))

    doc.build(story)


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX 生成（依赖 python-docx）
# ─────────────────────────────────────────────────────────────────────────────

def export_docx(config: ReportConfig, assets: ReportAssets, path: str):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        raise RuntimeError("DOCX 导出需要安装 python-docx：pip install python-docx")

    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
        section.top_margin    = Inches(1.1)
        section.bottom_margin = Inches(1.0)

    # 标题
    h = doc.add_heading(config.title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_txt = f"生成时间：{assets.timestamp}"
    if config.author:
        sub_txt += f"    作者：{config.author}"
    sub = doc.add_paragraph(sub_txt)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x80, 0xA0)

    if config.note:
        np = doc.add_paragraph(f"备注：{config.note}")
        np.runs[0].italic = True
        np.runs[0].font.size = Pt(10)

    # 元数据
    if config.include_meta:
        doc.add_heading("元数据信息", level=1)
        rows = [("报告时间", assets.timestamp), ("作者", config.author or "—")]
        rows += [(str(k), str(v)) for k, v in assets.metadata.items()]
        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(rows):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = v
            for cell in tbl.rows[i].cells:
                for para in cell.paragraphs:
                    para.runs[0].font.size = Pt(9) if para.runs else None

    # 地图截图
    if config.include_map and assets.map_png:
        doc.add_heading("地图视图", level=1)
        doc.add_picture(io.BytesIO(assets.map_png), width=Inches(5.5))
        cap = doc.add_paragraph("图：当前地图视图截图")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(0x66, 0x80, 0xA0)

    # 插件图表
    if config.include_figures:
        for i, fig in enumerate(assets.figures, 1):
            doc.add_heading(fig["title"], level=1)
            doc.add_picture(io.BytesIO(fig["png"]), width=Inches(5.5))
            cap = doc.add_paragraph(f"图 {i}：{fig['title']}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap.runs:
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(0x66, 0x80, 0xA0)

    # 页脚
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f"由极地海冰多参数综合反演平台自动生成 · {assets.timestamp}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

    doc.save(path)


# ─────────────────────────────────────────────────────────────────────────────
#  便利函数：检查依赖
# ─────────────────────────────────────────────────────────────────────────────

def check_deps() -> dict:
    """返回各格式依赖的可用状态 {'html': True, 'pdf': bool, 'docx': bool}。"""
    result = {"html": True, "pdf": False, "docx": False}
    try:
        import reportlab  # noqa: F401
        result["pdf"] = True
    except ImportError:
        pass
    try:
        import docx  # noqa: F401
        result["docx"] = True
    except ImportError:
        pass
    return result
